#!/usr/bin/env python3
"""Czytanie konkretnych paragrafów z dokumentu DOCX."""

from docx import Document
import sys

def read_paragraph(docx_path, para_num):
    """Czyta konkretny paragraf z dokumentu."""
    doc = Document(docx_path)

    if para_num < 0 or para_num >= len(doc.paragraphs):
        print(f"Błąd: Paragraf #{para_num} nie istnieje. Dokument ma {len(doc.paragraphs)} paragrafów.")
        return

    para = doc.paragraphs[para_num]
    print(f"\n=== Paragraf #{para_num} ===")
    print(f"Styl: {para.style.name}")
    print(f"Długość: {len(para.text)} znaków")
    print(f"\nTreść:")
    print("-" * 80)
    print(para.text)
    print("-" * 80)

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Użycie: python read_paragraphs.py <plik.docx> <numer_paragrafu> [numer2] [numer3]...")
        sys.exit(1)

    docx_path = sys.argv[1]
    para_nums = [int(num) for num in sys.argv[2:]]

    for num in para_nums:
        read_paragraph(docx_path, num)
