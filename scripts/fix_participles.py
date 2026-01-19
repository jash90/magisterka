#!/usr/bin/env python3
"""Poprawia imiesłowy kończące zdania w dokumentach DOCX."""

from docx import Document
import re
import sys

def fix_paragraph_38(doc):
    """Poprawia paragraf #38 - zakazane słowo i imiesłów."""
    para = doc.paragraphs[38]
    text = para.text

    # Fix 1: "złożonych chorób" → "poważnych chorób"
    text = text.replace("złożonych chorób", "poważnych chorób")

    # Fix 2: Imiesłów kończący zdanie
    # PRZED: "...system zabezpieczeń (guardrails) zapewniający bezpieczeństwo medyczne i etyczne użytkowania systemu."
    # PO: "...system zabezpieczeń (guardrails), który zapewnia bezpieczeństwo medyczne i etyczne użytkowania."

    old_phrase = "system zabezpieczeń (guardrails) zapewniający bezpieczeństwo medyczne i etyczne użytkowania systemu"
    new_phrase = "system zabezpieczeń (guardrails), który zapewnia bezpieczeństwo medyczne i etyczne użytkowania"

    text = text.replace(old_phrase, new_phrase)

    para.text = text
    print(" Paragraf #38 poprawiony")

def fix_paragraph_90(doc):
    """Poprawia paragraf #90 - zakazane słowo i imiesłów."""
    para = doc.paragraphs[90]
    text = para.text

    # Fix 1: "złożoność i formę wyjaśnień" → "formę i szczegółowość wyjaśnień"
    text = text.replace("złożoność i formę wyjaśnień", "formę i szczegółowość wyjaśnień")

    # Fix 2: Imiesłów kończący punkt 5
    # PRZED: "5. Opracowanie systemu zabezpieczeń (guardrails) zapewniającego bezpieczne i etyczne użytkowanie systemu w kontekście medycznym."
    # PO: "5. Opracowanie systemu zabezpieczeń (guardrails), który zapewnia bezpieczne i etyczne użytkowanie systemu w kontekście medycznym."

    old_phrase = "Opracowanie systemu zabezpieczeń (guardrails) zapewniającego bezpieczne i etyczne użytkowanie systemu w kontekście medycznym."
    new_phrase = "Opracowanie systemu zabezpieczeń (guardrails), który zapewnia bezpieczne i etyczne użytkowanie systemu w kontekście medycznym."

    text = text.replace(old_phrase, new_phrase)

    para.text = text
    print(" Paragraf #90 poprawiony")

def fix_paragraph_166(doc):
    """Poprawia paragraf #166 - 2 imiesłowy kończące zdania."""
    para = doc.paragraphs[166]
    text = para.text

    # Fix 1: Imiesłów w pierwszym zdaniu
    # PRZED: "...założone cele badawcze, dostarczając funkcjonalny system..."
    # PO: "...założone cele badawcze. System dostarcza funkcjonalne rozwiązanie..."

    old_phrase_1 = "zrealizowała założone cele badawcze, dostarczając funkcjonalny system wyjaśnialnej sztucznej inteligencji do predykcji śmiertelności w zapaleniu naczyń."
    new_phrase_1 = "zrealizowała założone cele badawcze. Opracowany system wyjaśnialnej sztucznej inteligencji umożliwia predykcję śmiertelności w zapaleniu naczyń."

    text = text.replace(old_phrase_1, new_phrase_1)

    # Fix 2: Imiesłów w punkcie 5
    # PRZED: "5.  Zaimplementowano system guardrails zapewniający bezpieczeństwo medyczne"
    # PO: "5.  Zaimplementowano system guardrails, który zapewnia bezpieczeństwo medyczne"

    old_phrase_2 = "Zaimplementowano system guardrails zapewniający bezpieczeństwo medyczne"
    new_phrase_2 = "Zaimplementowano system guardrails, który zapewnia bezpieczeństwo medyczne"

    text = text.replace(old_phrase_2, new_phrase_2)

    para.text = text
    print(" Paragraf #166 poprawiony")

def fix_paragraph_100(doc):
    """Poprawia paragraf #100 - zakazane słowo w kontekście technicznym."""
    para = doc.paragraphs[100]
    text = para.text

    # W kontekście LIME, "złożony model" jest terminem technicznym
    # Ale według instrukcji, "złożony" jest zakazany, więc zmienię na "skomplikowany"
    text = text.replace("złożonego modelu", "skomplikowanego modelu")

    para.text = text
    print(" Paragraf #100 poprawiony")

def fix_document(input_path, output_path):
    """Główna funkcja poprawiająca dokument."""
    print(f"\n=== Poprawianie dokumentu: {input_path} ===\n")

    doc = Document(input_path)

    # Poprawki paragrafów
    fix_paragraph_38(doc)
    fix_paragraph_90(doc)
    fix_paragraph_166(doc)
    fix_paragraph_100(doc)

    # Zapisz poprawiony dokument
    doc.save(output_path)
    print(f"\n Dokument zapisany: {output_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Użycie: python fix_participles.py <plik_wejściowy.docx> <plik_wyjściowy.docx>")
        sys.exit(1)

    fix_document(sys.argv[1], sys.argv[2])
