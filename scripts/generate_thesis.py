#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generator dokumentu pracy magisterskiej - Semestr I
System XAI do predykcji śmiertelności w zapaleniu naczyń

Autor: Bartłomiej Zimny
Data: 12 stycznia 2026
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import os
from datetime import datetime


def create_thesis_document():
    """Główna funkcja generująca dokument pracy magisterskiej"""

    print(" Rozpoczynam generowanie dokumentu pracy magisterskiej...")

    doc = Document()

    # 1. Konfiguracja stylów
    print("    Konfiguracja stylów i marginesy...")
    setup_styles(doc)

    # 2. Strona tytułowa
    print("    Tworzenie strony tytułowej...")
    add_title_page(doc)

    # 3. Oświadczenie
    print("     Dodawanie oświadczenia autora...")
    add_declaration(doc)

    # 4. Streszczenie PL
    print("    Generowanie streszczenia (PL)...")
    add_abstract_pl(doc)

    # 5. Abstract EN
    print("    Generowanie abstract (EN)...")
    add_abstract_en(doc)

    # 6. Spisy
    print("    Tworzenie spisów...")
    add_table_of_contents(doc)
    add_list_of_figures(doc)
    add_list_of_tables(doc)
    add_abbreviations(doc)

    # 7. Rozdziały
    print("    Rozdział 1: Wstęp...")
    add_chapter_1_introduction(doc)

    print("    Rozdział 2: Przegląd literatury...")
    add_chapter_2_literature_review(doc)

    print("    Rozdział 3: Metodologia...")
    add_chapter_3_methodology(doc)

    print("    Rozdział 4: Implementacja...")
    add_chapter_4_implementation(doc)

    print("    Rozdział 5: Eksperymenty i wyniki...")
    add_chapter_5_results(doc)

    print("    Rozdział 6: Dyskusja...")
    add_chapter_6_discussion(doc)

    print("    Rozdział 7: Wnioski...")
    add_chapter_7_conclusions(doc)

    # 8. Bibliografia
    print("    Generowanie bibliografii...")
    add_bibliography(doc)

    # 9. Załączniki
    print("    Dodawanie załączników...")
    add_appendices(doc)

    # 10. Zapisanie
    output_path = 'docs/magisterka_Zimny_sem1.docx'
    os.makedirs('docs', exist_ok=True)
    doc.save(output_path)

    print(f"\n Dokument pracy magisterskiej wygenerowany!")
    print(f"   Lokalizacja: {output_path}")
    print(f"   Rozmiar: {os.path.getsize(output_path) / 1024:.1f} KB")

    return output_path


def setup_styles(doc):
    """Konfiguracja stylów dokumentu zgodnie z polskimi standardami"""

    # Marginesy: lewy 3.5cm, prawy 1.5cm, górny/dolny 2.5cm
    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21.0)   # A4

    # Style nagłówków
    styles = doc.styles

    # Heading 1 - Rozdziały główne
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(12)
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2 - Podrozdziały
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(13)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Pt(18)
    h2_style.paragraph_format.space_after = Pt(6)

    # Heading 3 - Podpodrozdziały
    h3_style = styles['Heading 3']
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(6)

    # Normal text - Tekst główny
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.space_after = Pt(6)


def add_title_page(doc):
    """Strona tytułowa zgodna z polskimi standardami akademickimi"""

    # Nazwa uczelni (centered, większa czcionka)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AKADEMIA GÓRNICZO-HUTNICZA")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IM. STANISŁAWA STASZICA W KRAKOWIE")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()  # Spacer

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Wydział Informatyki, Elektroniki i Telekomunikacji")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Kierunek: Informatyka")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Spacer
    for _ in range(5):
        doc.add_paragraph()

    # Tytuł pracy
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PRACA MAGISTERSKA")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()  # Spacer

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("System wyjaśnialnej sztucznej inteligencji (XAI)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("do predykcji śmiertelności w zapaleniu naczyń")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Autor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Autor:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bartłomiej Zimny")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_paragraph()  # Spacer

    # Promotor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Promotor:")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Tytuł i imię nazwisko promotora]")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    # Spacer
    for _ in range(3):
        doc.add_paragraph()

    # Miejsce i rok
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Kraków 2026")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Page break
    doc.add_page_break()


def add_declaration(doc):
    """Oświadczenie autora"""

    doc.add_heading('OŚWIADCZENIE AUTORA', level=1)

    text = """Oświadczam, że niniejsza praca została przeze mnie wykonana samodzielnie i nie zawiera treści uzyskanych w sposób niezgodny z obowiązującymi przepisami.

Oświadczam również, że przedstawiona praca nie była wcześniej przedmiotem procedur związanych z uzyskaniem dyplomu wyższej uczelni.

Oświadczam ponadto, że niniejsza wersja pracy jest identyczna z załączoną wersją elektroniczną."""

    doc.add_paragraph(text)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph("Data: .................................")
    p.add_run("\t\t").add_text("Podpis: .................................")

    doc.add_page_break()


def add_abstract_pl(doc):
    """Streszczenie w języku polskim"""

    doc.add_heading('STRESZCZENIE', level=1)

    abstract = """Zapalenie naczyń stanowi grupę złożonych chorób autoimmunologicznych charakteryzujących się wysoką śmiertelnością i trudnością w prognozowaniu przebiegu klinicznego. Niniejsza praca magisterska przedstawia kompleksowy system wyjaśnialnej sztucznej inteligencji (XAI) przeznaczony do predykcji ryzyka śmiertelności u pacjentów z zapaleniem naczyń.

Głównym celem pracy jest opracowanie, implementacja i ewaluacja systemu uczenia maszynowego, który nie tylko dokładnie przewiduje ryzyko śmiertelności, ale również dostarcza przejrzystych, zrozumiałych wyjaśnień decyzji modelu, dostosowanych do potrzeb różnych grup odbiorców – od klinicystów po pacjentów.

W ramach pracy zaimplementowano i porównano siedem algorytmów uczenia maszynowego (Random Forest, XGBoost, LightGBM, MLP, Regresja Logistyczna, SVM, Gradient Boosting), przy czym szczególną uwagę poświęcono optymalizacji metryk medycznych, w szczególności czułości (sensitivity), kluczowej dla wykrywania przypadków zagrożenia życia.

Integralną częścią systemu jest moduł wyjaśnialnej sztucznej inteligencji, który łączy cztery komplementarne metody XAI: LIME (Local Interpretable Model-agnostic Explanations), SHAP (SHapley Additive exPlanations), DALEX (Descriptive mAchine Learning EXplanations) oraz EBM (Explainable Boosting Machine). Porównanie tych metod pozwoliło na identyfikację kluczowych czynników prognostycznych oraz ocenę zgodności wyjaśnień generowanych przez różne podejścia.

System zrealizowano jako kompletną aplikację webową składającą się z backendu (API RESTful oparte na FastAPI), frontendu (interaktywny dashboard w Streamlit) oraz agenta konwersacyjnego wykorzystującego RAG (Retrieval-Augmented Generation) i duże modele językowe. Dodatkowo zaimplementowano system zabezpieczeń (guardrails) zapewniający bezpieczeństwo medyczne i etyczne użytkowania systemu.

Wyniki eksperymentalne wykazały wysoką skuteczność modelu XGBoost w predykcji śmiertelności, przy jednoczesnym zachowaniu interpretowalności dzięki metodom XAI. Analiza porównawcza metod XAI wykazała wysoką zgodność w identyfikacji kluczowych cech predykcyjnych, takich jak wiek pacjenta, zajęcie nerek oraz liczba zajętych narządów.

System stanowi innowacyjne rozwiązanie łączące zaawansowane uczenie maszynowe z wymogami transparentności w medycynie, oferując wsparcie dla decyzji klinicznych przy zachowaniu wysokich standardów etycznych i bezpieczeństwa."""

    doc.add_paragraph(abstract)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Słowa kluczowe: ")
    run.font.bold = True
    p.add_run("XAI, SHAP, LIME, uczenie maszynowe, medycyna, predykcja śmiertelności, zapalenie naczyń, wyjaśnialna sztuczna inteligencja, DALEX, EBM")

    doc.add_page_break()


def add_abstract_en(doc):
    """Abstract in English"""

    doc.add_heading('ABSTRACT', level=1)

    abstract = """Vasculitis represents a group of complex autoimmune diseases characterized by high mortality rates and difficulty in predicting clinical outcomes. This master's thesis presents a comprehensive explainable artificial intelligence (XAI) system designed to predict mortality risk in patients with vasculitis.

The main objective of this work is to develop, implement, and evaluate a machine learning system that not only accurately predicts mortality risk but also provides transparent, understandable explanations of model decisions, tailored to the needs of different stakeholder groups – from clinicians to patients.

Seven machine learning algorithms were implemented and compared (Random Forest, XGBoost, LightGBM, MLP, Logistic Regression, SVM, Gradient Boosting), with particular attention paid to optimizing medical metrics, especially sensitivity, which is crucial for detecting life-threatening cases.

An integral part of the system is the explainable AI module, which combines four complementary XAI methods: LIME (Local Interpretable Model-agnostic Explanations), SHAP (SHapley Additive exPlanations), DALEX (Descriptive mAchine Learning EXplanations), and EBM (Explainable Boosting Machine). Comparison of these methods enabled identification of key prognostic factors and assessment of explanation agreement across different approaches.

The system was implemented as a complete web application consisting of a backend (RESTful API based on FastAPI), frontend (interactive Streamlit dashboard), and a conversational agent utilizing RAG (Retrieval-Augmented Generation) and large language models. Additionally, a guardrail system was implemented to ensure medical and ethical safety of system usage.

Experimental results demonstrated high effectiveness of the XGBoost model in mortality prediction while maintaining interpretability through XAI methods. Comparative analysis of XAI methods showed high agreement in identifying key predictive features such as patient age, kidney involvement, and number of affected organs.

The system represents an innovative solution combining advanced machine learning with transparency requirements in medicine, offering support for clinical decisions while maintaining high ethical and safety standards."""

    doc.add_paragraph(abstract)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.font.bold = True
    p.add_run("XAI, SHAP, LIME, machine learning, medicine, mortality prediction, vasculitis, explainable artificial intelligence, DALEX, EBM")

    doc.add_page_break()


def add_table_of_contents(doc):
    """Spis treści - placeholder dla automatycznego generowania w Word"""

    doc.add_heading('SPIS TREŚCI', level=1)

    p = doc.add_paragraph()
    run = p.add_run("[Spis treści zostanie wygenerowany automatycznie w Microsoft Word]")
    run.font.italic = True

    p = doc.add_paragraph()
    run = p.add_run("Instrukcja: ")
    run.font.bold = True
    p.add_run("W Microsoft Word: Referencje → Spis treści → Automatyczny spis treści")

    doc.add_page_break()


def add_list_of_figures(doc):
    """Spis rysunków"""

    doc.add_heading('SPIS RYSUNKÓW', level=1)

    figures = [
        "Rysunek 1.1. Architektura systemu Vasculitis XAI",
        "Rysunek 2.1. Taksonomia metod XAI",
        "Rysunek 2.2. Schemat działania algorytmu LIME",
        "Rysunek 2.3. Wykres waterfall SHAP - przykład wyjaśnienia",
        "Rysunek 3.1. Pipeline przetwarzania danych",
        "Rysunek 3.2. Schemat walidacji krzyżowej",
        "Rysunek 4.1. Diagram architektury API",
        "Rysunek 4.2. Screenshot dashboardu - formularz wejściowy",
        "Rysunek 4.3. Screenshot dashboardu - wizualizacja wyników",
        "Rysunek 5.1. Krzywa ROC-AUC dla modelu XGBoost",
        "Rysunek 5.2. Wykres feature importance - top 10 cech",
        "Rysunek 5.3. Wyjaśnienie SHAP - przykładowy pacjent",
        "Rysunek 5.4. Wyjaśnienie LIME - przykładowy pacjent",
        "Rysunek 5.5. Porównanie metod XAI - zgodność rankingów",
    ]

    for figure in figures:
        doc.add_paragraph(figure, style='List Number')

    doc.add_page_break()


def add_list_of_tables(doc):
    """Spis tabel"""

    doc.add_heading('SPIS TABEL', level=1)

    tables = [
        "Tabela 2.1. Porównanie metod XAI - charakterystyka",
        "Tabela 3.1. Charakterystyka zbioru danych",
        "Tabela 3.2. Zmienne wejściowe (cechy kliniczne)",
        "Tabela 3.3. Hiperparametry modeli uczenia maszynowego",
        "Tabela 3.4. Strategie radzenia sobie z niezbalansowaniem klas",
        "Tabela 3.5. Metryki medyczne - progi akceptacji",
        "Tabela 5.1. Porównanie wydajności modeli ML",
        "Tabela 5.2. Top 10 cech predykcyjnych - feature importance",
        "Tabela 5.3. Zgodność metod XAI - korelacja Spearmana",
        "Tabela 5.4. Czas generowania wyjaśnień dla różnych metod XAI",
        "Tabela 5.5. Wydajność API - czas odpowiedzi endpointów",
    ]

    for table in tables:
        doc.add_paragraph(table, style='List Number')

    doc.add_page_break()


def add_abbreviations(doc):
    """Wykaz skrótów i symboli"""

    doc.add_heading('WYKAZ SKRÓTÓW I SYMBOLI', level=1)

    abbreviations = [
        ("AI", "Artificial Intelligence (Sztuczna Inteligencja)"),
        ("API", "Application Programming Interface"),
        ("AUC-ROC", "Area Under the Receiver Operating Characteristic Curve"),
        ("CDSS", "Clinical Decision Support System (System Wsparcia Decyzji Klinicznych)"),
        ("CSN", "Centralny System Nerwowy"),
        ("DALEX", "Descriptive mAchine Learning EXplanations"),
        ("EBM", "Explainable Boosting Machine"),
        ("EHR", "Electronic Health Record (Elektroniczna Dokumentacja Medyczna)"),
        ("EULAR", "European League Against Rheumatism"),
        ("GAM", "Generalized Additive Model"),
        ("GDPR", "General Data Protection Regulation"),
        ("LIME", "Local Interpretable Model-agnostic Explanations"),
        ("LLM", "Large Language Model (Duży Model Językowy)"),
        ("MCC", "Matthews Correlation Coefficient"),
        ("ML", "Machine Learning (Uczenie Maszynowe)"),
        ("MLP", "Multi-Layer Perceptron (Wielowarstwowy Perceptron)"),
        ("NPV", "Negative Predictive Value (Ujemna Wartość Predykcyjna)"),
        ("OIT", "Oddział Intensywnej Terapii"),
        ("PPV", "Positive Predictive Value (Dodatnia Wartość Predykcyjna)"),
        ("RAG", "Retrieval-Augmented Generation"),
        ("REST", "Representational State Transfer"),
        ("SHAP", "SHapley Additive exPlanations"),
        ("SMOTE", "Synthetic Minority Oversampling Technique"),
        ("SVM", "Support Vector Machine (Maszyna Wektorów Wspierających)"),
        ("XAI", "Explainable Artificial Intelligence (Wyjaśnialna Sztuczna Inteligencja)"),
        ("XGBoost", "eXtreme Gradient Boosting"),
    ]

    table = doc.add_table(rows=len(abbreviations), cols=2)
    table.style = 'Light Grid Accent 1'

    for idx, (abbr, definition) in enumerate(abbreviations):
        row = table.rows[idx]
        row.cells[0].text = abbr
        row.cells[1].text = definition

        # Formatowanie
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles['Normal']
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    doc.add_page_break()


def add_chapter_1_introduction(doc):
    """Rozdział 1: Wstęp"""

    doc.add_heading('1. WSTĘP', level=1)

    # 1.1. Kontekst i motywacja
    doc.add_heading('1.1. Kontekst i motywacja badań', level=2)

    intro_1_1 = """Zapalenie naczyń (vasculitis) to grupa heterogenicznych chorób autoimmunologicznych charakteryzujących się zapaleniem i uszkodzeniem ścian naczyń krwionośnych różnej wielkości. Pomimo postępów w diagnostyce i leczeniu, śmiertelność w zapaleniu naczyń pozostaje istotnym problemem klinicznym, sięgając w niektórych postaciach nawet 20-30% w ciągu pierwszych pięciu lat od rozpoznania.

Precyzyjna ocena ryzyka śmiertelności u pacjentów z zapaleniem naczyń ma kluczowe znaczenie dla podejmowania decyzji terapeutycznych, alokacji zasobów medycznych oraz informowania pacjentów o rokowaniu. Tradycyjne skale prognostyczne opierają się na ograniczonej liczbie parametrów klinicznych i często nie wykorzystują pełnego potencjału dostępnych danych medycznych.

Rozwój metod uczenia maszynowego (ML) otwiera nowe możliwości w zakresie predykcji medycznej, oferując modele o wyższej dokładności niż tradycyjne podejścia statystyczne. Jednakże zastosowanie algorytmów ML w medycynie napotyka na fundamentalną barierę – brak transparentności i interpretowalności tzw. "czarnych skrzynek". W kontekście klinicznym, gdzie decyzje mogą mieć bezpośredni wpływ na życie pacjenta, zrozumienie podstaw decyzji modelu nie jest opcjonalne, lecz wymogiem prawnym, etycznym i praktycznym.

Wyjaśnialna sztuczna inteligencja (Explainable AI, XAI) stanowi odpowiedź na to wyzwanie. Metody XAI umożliwiają zrozumienie, dlaczego model dokonał określonej predykcji, identyfikują kluczowe czynniki wpływające na wynik oraz pozwalają na walidację medycznego sensu decyzji algorytmu. W ostatnich latach rozwinięto szereg metod XAI, w tym LIME, SHAP, DALEX oraz inherently interpretable models jak EBM, każda z własnymi zaletami i ograniczeniami.

Motywacją do podjęcia niniejszych badań była obserwacja luki pomiędzy możliwościami technicznymi ML a wymogami praktyki klinicznej. Istniejące systemy predykcyjne często skupiają się wyłącznie na maksymalizacji dokładności, zaniedbując aspekty interpretowalności, komunikacji z użytkownikiem końcowym oraz bezpieczeństwa medycznego. Brakuje również kompleksowych porównań różnych metod XAI w kontekście konkretnych zastosowań medycznych."""

    doc.add_paragraph(intro_1_1)

    # 1.2. Cele pracy
    doc.add_heading('1.2. Cele pracy', level=2)

    intro_1_2 = """Głównym celem niniejszej pracy magisterskiej jest opracowanie i ewaluacja systemu wyjaśnialnej sztucznej inteligencji do predykcji śmiertelności u pacjentów z zapaleniem naczyń, który łączy wysoką dokładność predykcyjną z przejrzystością i użytecznością kliniczną.

Szczegółowe cele badawcze obejmują:

1. Implementację i porównanie wydajności różnych algorytmów uczenia maszynowego (Random Forest, XGBoost, LightGBM, MLP, Regresja Logistyczna, SVM, Gradient Boosting) w kontekście predykcji śmiertelności w zapaleniu naczyń, z uwzględnieniem medycznych metryk ewaluacji.

2. Opracowanie i integrację czterech komplementarnych metod wyjaśnialnej sztucznej inteligencji (LIME, SHAP, DALEX, EBM) w celu generowania lokalnych i globalnych wyjaśnień predykcji modelu.

3. Przeprowadzenie systematycznego porównania metod XAI pod kątem zgodności generowanych wyjaśnień, stabilności, czasu obliczeń oraz użyteczności w kontekście medycznym.

4. Zaprojektowanie i implementację kompletnego systemu wsparcia decyzji klinicznych, obejmującego backend API, interaktywny dashboard oraz agenta konwersacyjnego.

5. Opracowanie systemu zabezpieczeń (guardrails) zapewniającego bezpieczne i etyczne użytkowanie systemu w kontekście medycznym.

6. Implementację mechanizmów komunikacji wielopoziomowej, dostosowujących złożoność i formę wyjaśnień do poziomu wiedzy medycznej odbiorcy (klinicysta, pacjent zaawansowany, pacjent podstawowy).

7. Ewaluację praktycznej użyteczności systemu w kontekście rzeczywistych scenariuszy klinicznych oraz identyfikację kierunków przyszłego rozwoju."""

    doc.add_paragraph(intro_1_2)

    # 1.3. Hipotezy badawcze
    doc.add_heading('1.3. Hipotezy badawcze', level=2)

    intro_1_3 = """W ramach niniejszej pracy sformułowano następujące hipotezy badawcze:

**H1:** Modele uczenia maszynowego oparte na gradient boosting (XGBoost, LightGBM) osiągną wyższą skuteczność w predykcji śmiertelności w zapaleniu naczyń niż tradycyjne metody statystyczne (regresja logistyczna), mierzoną jako AUC-ROC ≥ 0.75 oraz czułość (sensitivity) ≥ 0.80.

**H2:** Integracja wielu metod XAI (LIME, SHAP, DALEX, EBM) zapewni wyższą wiarygodność i stabilność wyjaśnień niż pojedyncza metoda, poprzez identyfikację konsensusowych cech predykcyjnych.

**H3:** Różne metody XAI wykażą wysoką zgodność (korelacja Spearmana ≥ 0.70) w rankingu najważniejszych cech predykcyjnych, co potwierdzi ich wiarygodność w kontekście klinicznym.

**H4:** System wyjaśnialnej sztucznej inteligencji z mechanizmami dostosowania komunikacji do poziomu wiedzy medycznej użytkownika zwiększy użyteczność i akceptowalność systemu ML w praktyce klinicznej.

**H5:** Implementacja guardrails (zabezpieczeń medycznych i etycznych) w systemie XAI jest niezbędna i możliwa do realizacji bez znaczącego wpływu na funkcjonalność predykcyjną systemu."""

    doc.add_paragraph(intro_1_3)

    # 1.4. Zakres pracy
    doc.add_heading('1.4. Zakres pracy', level=2)

    intro_1_4 = """Zakres niniejszej pracy magisterskiej obejmuje następujące obszary:

**Zakres teoretyczny:**
- Przegląd literatury z zakresu uczenia maszynowego w medycynie
- Analiza metod wyjaśnialnej sztucznej inteligencji (LIME, SHAP, DALEX, EBM)
- Studium metryk ewaluacji w kontekście medycznym
- Analiza systemów wsparcia decyzji klinicznych (CDSS)
- Przegląd aspektów etycznych i regulacyjnych AI w medycynie

**Zakres eksperymentalny:**
- Przetwarzanie i przygotowanie danych klinicznych pacjentów z zapaleniem naczyń
- Implementacja i trenowanie siedmiu algorytmów uczenia maszynowego
- Optymalizacja hiperparametrów z uwzględnieniem metryk medycznych
- Implementacja czterech metod XAI (LIME, SHAP, DALEX, EBM)
- Opracowanie framework'u do porównania metod XAI
- Ewaluacja wydajności modeli i jakości wyjaśnień

**Zakres implementacyjny:**
- Architektura i implementacja API RESTful (FastAPI)
- Implementacja interaktywnego dashboardu (Streamlit)
- Integracja agenta konwersacyjnego (LangChain + RAG)
- System zabezpieczeń medycznych (guardrails)
- Konteneryzacja i deployment (Docker)
- Testy jednostkowe i integracyjne (pytest)

**Poza zakresem pracy:**
- Walidacja prospektywna w warunkach rzeczywistych
- Certyfikacja medyczna systemu (Medical Device Regulation)
- Integracja z rzeczywistymi systemami szpitalnymi (EHR/EMR)
- Badania kliniczne z udziałem pacjentów"""

    doc.add_paragraph(intro_1_4)

    # 1.5. Struktura pracy
    doc.add_heading('1.5. Struktura pracy', level=2)

    intro_1_5 = """Niniejsza praca została podzielona na siedem rozdziałów:

**Rozdział 1 - Wstęp** przedstawia kontekst i motywację badań, formułuje cele i hipotezy badawcze oraz definiuje zakres pracy.

**Rozdział 2 - Przegląd literatury i podstawy teoretyczne** zawiera szczegółową analizę stanu wiedzy w zakresie uczenia maszynowego w medycynie, metod wyjaśnialnej sztucznej inteligencji (LIME, SHAP, DALEX, EBM), metryk medycznych oraz systemów wsparcia decyzji klinicznych.

**Rozdział 3 - Metodologia** opisuje metodykę badań, charakterystykę zbioru danych, pipeline przetwarzania danych, algorytmy uczenia maszynowego, metody XAI oraz kryteria ewaluacji systemu.

**Rozdział 4 - Implementacja** przedstawia architekturę systemu oraz szczegóły implementacyjne poszczególnych modułów: przetwarzania danych, trenowania modeli, generowania wyjaśnień XAI, API RESTful, dashboardu oraz systemu zabezpieczeń.

**Rozdział 5 - Eksperymenty i wyniki** prezentuje wyniki eksperymentów, w tym porównanie wydajności modeli ML, analizę ważności cech, wyniki metod XAI oraz ich porównanie, a także ewaluację wydajności systemu.

**Rozdział 6 - Dyskusja** zawiera interpretację uzyskanych wyników, porównanie metod XAI, dyskusję znaczenia dla praktyki klinicznej, aspekty etyczne i prawne, ograniczenia badania oraz porównanie z pracami pokrewnymi.

**Rozdział 7 - Wnioski i prace przyszłe** podsumowuje osiągnięcia pracy, prezentuje główne wnioski oraz wskazuje kierunki przyszłych badań i rozszerzeń systemu."""

    doc.add_paragraph(intro_1_5)

    doc.add_page_break()


# Kontynuacja skryptu w następnej części...

if __name__ == "__main__":
    try:
        output_file = create_thesis_document()
        print(f"\n✨ Sukces! Plik zapisany: {output_file}")
        print("\n Następne kroki:")
        print("   1. Otwórz dokument w Microsoft Word")
        print("   2. Wygeneruj automatyczny spis treści")
        print("   3. Dodaj wykresy i diagramy")
        print("   4. Uzupełnij dane promotora na stronie tytułowej")
        print("   5. Dokonaj końcowej korekty")

    except Exception as e:
        print(f"\n Błąd podczas generowania dokumentu: {e}")
        import traceback
        traceback.print_exc()
