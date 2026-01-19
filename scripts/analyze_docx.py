#!/usr/bin/env python3
"""Analiza dokumentu DOCX dla znalezienia cech AI do usunięcia."""

import re
from docx import Document

def analyze_document(docx_path):
    """Analizuje dokument DOCX i znajduje problemy."""
    doc = Document(docx_path)

    problems = {
        'participles': [],  # Imiesłowy kończące zdania
        'forbidden_words': [],  # Zakazane słowa
        'title_case_headers': []  # Nagłówki w Title Case
    }

    print(f"\n=== Analiza dokumentu: {docx_path} ===\n")
    print(f"Liczba paragrafów: {len(doc.paragraphs)}\n")

    # Wzorce do wyszukiwania
    participle_pattern = re.compile(r'\b(zapewniając|zapewniający|zapewniającego|umożliwiając|umożliwiający)\b', re.IGNORECASE)
    forbidden_pattern = re.compile(r'\b(złożon\w*)\b', re.IGNORECASE)

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if not text:
            continue

        # Sprawdzenie imiesłowów kończących zdania
        if participle_pattern.search(text):
            # Sprawdź czy kończy zdanie (ostatnie słowo przed kropką)
            sentences = text.split('.')
            for sent_idx, sentence in enumerate(sentences):
                if participle_pattern.search(sentence.strip()[-30:]):  # Ostatnie 30 znaków zdania
                    problems['participles'].append({
                        'para_num': idx,
                        'text': text[:100] + '...' if len(text) > 100 else text,
                        'full_text': text
                    })
                    break

        # Sprawdzenie zakazanych słów
        if forbidden_pattern.search(text):
            problems['forbidden_words'].append({
                'para_num': idx,
                'text': text[:100] + '...' if len(text) > 100 else text,
                'full_text': text,
                'match': forbidden_pattern.search(text).group()
            })

        # Sprawdzenie nagłówków Title Case
        if para.style.name.startswith('Heading') or (text and len(text) < 100):
            words = text.split()
            if len(words) >= 2:
                # Sprawdź czy więcej niż 1 słowo zaczyna się z wielkiej litery (oprócz akronimów)
                capital_words = [w for w in words if w and w[0].isupper() and len(w) > 1 and not w.isupper()]
                if len(capital_words) >= 2:
                    problems['title_case_headers'].append({
                        'para_num': idx,
                        'text': text,
                        'style': para.style.name
                    })

    # Wyświetlenie wyników
    print("=" * 80)
    print(f"IMIESŁOWY KOŃCZĄCE ZDANIA: {len(problems['participles'])} znalezionych")
    print("=" * 80)
    for p in problems['participles']:
        print(f"\nParagraf #{p['para_num']}:")
        print(f"  Tekst: {p['text']}")

    print("\n" + "=" * 80)
    print(f"ZAKAZANE SŁOWA: {len(problems['forbidden_words'])} znalezionych")
    print("=" * 80)
    for p in problems['forbidden_words']:
        print(f"\nParagraf #{p['para_num']}:")
        print(f"  Słowo: '{p['match']}'")
        print(f"  Tekst: {p['text']}")

    print("\n" + "=" * 80)
    print(f"NAGŁÓWKI TITLE CASE: {len(problems['title_case_headers'])} znalezionych")
    print("=" * 80)
    for p in problems['title_case_headers']:
        print(f"\nParagraf #{p['para_num']} ({p['style']}):")
        print(f"  Tekst: {p['text']}")

    return problems

if __name__ == '__main__':
    import sys

    if len(sys.argv) < 2:
        print("Użycie: python analyze_docx.py <ścieżka_do_pliku.docx>")
        sys.exit(1)

    analyze_document(sys.argv[1])
