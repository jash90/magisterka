#!/usr/bin/env python3
"""Usuwa zaawansowane cechy AI zgodnie z rozszerzonymi instrukcjami."""

from docx import Document
import sys

def fix_paragraph_88(doc):
    """Usuwa puste frazy 'istotny/istotne' z paragrafu #88."""
    para = doc.paragraphs[88]
    text = para.text

    # Fix 1: "pozostaje istotnym problemem" → "pozostaje poważnym problemem"
    text = text.replace("pozostaje istotnym problemem klinicznym",
                       "pozostaje poważnym problemem klinicznym")

    # Fix 2: "ma istotne znaczenie dla" → "wpływa na"
    text = text.replace("ma istotne znaczenie dla podejmowania decyzji terapeutycznych",
                       "wpływa na podejmowanie decyzji terapeutycznych")

    # Fix 3: "istotne czynniki" → "czynniki" (w kontekście XAI)
    text = text.replace("identyfikują istotne czynniki wpływające na wynik",
                       "identyfikują czynniki wpływające na wynik")

    para.text = text
    print("✓ Paragraf #88: usunięto puste frazy 'istotny/istotne'")

def fix_paragraph_97(doc):
    """Usuwa pustą frazę 'Istotne wyzwania' z paragrafu #97."""
    para = doc.paragraphs[97]
    text = para.text

    # Fix: "Istotne wyzwania obejmują" → "Wyzwania obejmują"
    text = text.replace("Istotne wyzwania obejmują", "Wyzwania obejmują")

    para.text = text
    print("✓ Paragraf #97: usunięto 'Istotne'")

def fix_paragraph_140(doc):
    """Usuwa pustą frazę 'jest istotna w kontekście' z paragrafu #140."""
    para = doc.paragraphs[140]
    text = para.text

    # Fix: "jest istotna w kontekście" → "w kontekście"
    text = text.replace("Wysoka czułość jest istotna w kontekście predykcji śmiertelności",
                       "Wysoka czułość w kontekście predykcji śmiertelności")

    para.text = text
    print("✓ Paragraf #140: usunięto pustą frazę 'jest istotna'")

def fix_paragraph_144(doc):
    """Usuwa problematyczne 'podkreślają' i 'odzwierciedlają' z paragrafu #144."""
    para = doc.paragraphs[144]
    text = para.text

    # Fix 1: "podkreślają krytyczne znaczenie" → "wskazują na znaczenie"
    # Nieożywione podmioty nie mogą "podkreślać"
    text = text.replace("razem podkreślają krytyczne znaczenie funkcji nerek",
                       "razem wskazują na znaczenie funkcji nerek")

    # Fix 2: "odzwierciedlająca ogólną prawidłowość" → "zgodna z ogólną prawidłowością"
    text = text.replace("odzwierciedlająca ogólną prawidłowość medyczną",
                       "zgodna z ogólną prawidłowością medyczną")

    # Fix 3: "odzwierciedlają aktywność" → "wskazują na aktywność"
    text = text.replace("odzwierciedlają aktywność choroby",
                       "wskazują na aktywność choroby")

    para.text = text
    print("✓ Paragraf #144: usunięto 'podkreślają' i 'odzwierciedlają'")

def fix_paragraph_156(doc):
    """Usuwa pustą frazę 'Szczególnie istotna jest' z paragrafu #156."""
    para = doc.paragraphs[156]
    text = para.text

    # Fix: "Szczególnie istotna jest wysoka czułość" → "Wysoka czułość"
    text = text.replace("Szczególnie istotna jest wysoka czułość (0.85), oznaczająca",
                       "Wysoka czułość (0.85) oznacza")

    para.text = text
    print("✓ Paragraf #156: usunięto 'Szczególnie istotna jest'")

def fix_paragraph_165_header(doc):
    """Zmienia nagłówek 'Podsumowanie osiągnięć' na bardziej konkretny."""
    para = doc.paragraphs[165]

    # Zmiana: "7.1. Podsumowanie osiągnięć" → "7.1. Realizacja celów badawczych"
    # Unikamy sekcji podsumowujących
    para.text = "7.1. Realizacja celów badawczych"

    print("✓ Paragraf #165: zmieniono nagłówek z 'Podsumowanie' na 'Realizacja'")

def fix_document(input_path, output_path):
    """Główna funkcja naprawiająca dokument."""
    print(f"\n=== Usuwanie zaawansowanych cech AI: {input_path} ===\n")

    doc = Document(input_path)

    # Poprawki paragrafów
    fix_paragraph_88(doc)
    fix_paragraph_97(doc)
    fix_paragraph_140(doc)
    fix_paragraph_144(doc)
    fix_paragraph_156(doc)
    fix_paragraph_165_header(doc)

    # Zapisz poprawiony dokument
    doc.save(output_path)
    print(f"\n✓ Dokument zapisany: {output_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Użycie: python fix_ai_phrases.py <plik_wejściowy.docx> <plik_wyjściowy.docx>")
        sys.exit(1)

    fix_document(sys.argv[1], sys.argv[2])
