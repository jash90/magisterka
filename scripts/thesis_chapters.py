#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Rozdziały pracy magisterskiej - funkcje generujące treść
Część 2: Rozdziały 2-7, Bibliografia, Załączniki
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_chapter_2_literature_review(doc):
    """Rozdział 2: Przegląd literatury i podstawy teoretyczne"""

    doc.add_heading('2. PRZEGLĄD LITERATURY I PODSTAWY TEORETYCZNE', level=1)

    # 2.1. Zapalenie naczyń
    doc.add_heading('2.1. Zapalenie naczyń - charakterystyka medyczna', level=2)

    content_2_1 = """Zapalenie naczyń (vasculitis) to heterogenna grupa chorób charakteryzujących się zapaleniem i uszkodzeniem ścian naczyń krwionośnych. Według klasyfikacji Chapel Hill Consensus Conference (CHCC) z 2012 roku, zapale naczyń dzieli się na kategorie w zależności od wielkości zajętych naczyń: zapale naczyń dużych (np. arteritis Takayasu, olbrzymiokomórkowe zapalenie tętnic), zapale naczyń średnich (np. guzkowe zapalenie tętnic, choroba Kawasaki) oraz zapalenia naczyń małych naczyń, które z kolei dzielą się na zapale związane z przeciwciałami ANCA i niezwiązane z ANCA.

Epidemiologia zapalenia naczyń jest zróżnicowana geograficznie i demograficznie. Zachorowalność waha się od 10 do 40 przypadków na milion mieszkańców rocznie, w zależności od typu zapalenia naczyń i populacji. Śmiertelność, pomimo postępów w leczeniu, pozostaje istotnym problem klinicznym, szczególnie w przypadkach z zajęciem nerek, układu oddechowego lub sercowo-naczyniowego.

Do kluczowych czynników ryzyka śmiertelności w zapaleniu naczyń należą:
- Starszy wiek w momencie rozpoznania
- Zajęcie nerek manifestujące się niewydolnością nerek
- Zajęcie układu sercowo-naczyniowego
- Zajęcie płuc z krwotokiem pęcherzykowym
- Zajęcie ośrodkowego układu nerwowego
- Liczba zajętych układów/narządów
- Opóźnienie w rozpoznaniu i rozpoczęciu leczenia
- Powikłania infekcyjne
- Nieodpowiedź na leczenie immunosupresyjne

Wyzwania w predykcji klinicznej wynikają z heterogenności przebiegu choroby, złożonych interakcji pomiędzy czynnikami ryzyka, względnie niskiej częstości występowania oraz indywidualnej zmienności odpowiedzi na leczenie."""

    doc.add_paragraph(content_2_1)

    # 2.2. Uczenie maszynowe w medycynie
    doc.add_heading('2.2. Uczenie maszynowe w medycynie', level=2)

    content_2_2 = """Uczenie maszynowe (Machine Learning, ML) znajduje rosnące zastosowanie w medycynie, obejmując diagnostykę, prognozowanie, personalizację terapii oraz odkrywanie nowych biomarkerów. W ostatniej dekadzie odnotowano przełomowe osiągnięcia w zastosowaniu głębokich sieci neuronowych do analizy obrazów medycznych, osiągających dokładność porównywalną lub wyższą od lekarzy specjalistów w zadaniach takich jak wykrywanie nowotworów skóry, rozpoznawanie retinopatii cukrzycowej czy klasyfikacja zmian w obrazach radiologicznych.

W zakresie modeli predykcyjnych, algorytmy ML wykazują przewagę nad tradycyjnymi skalami prognostycznymi w wielu obszarach medycyny. Chen i in. (2019) wykazali wyższą dokładność modeli gradient boosting w porównaniu z tradycyjnymi scorami klinicznymi w predykcji śmiertelności w sepsiie. Podobne obserwacje poczyniono w onkologii, kardiologii oraz nefologii.

Jednakże zastosowanie ML w medycynie wiąże się z istotnymi wyzwaniami etycznymi i regulacyjnymi. Główne obszary problemowe obejmują:

1. **Brak interpretowalności** - algorytmy ML działają często jako "czarne skrzynki", co utrudnia zrozumienie podstaw decyzji i weryfikację medycznego sensu predykcji.

2. **Bias i fairness** - modele ML mogą utrwalać lub wzmacniać istniejące nierówności w opiece zdrowotnej, jeśli dane treningowe nie są reprezentatywne.

3. **Odpowiedzialność prawna** - w przypadku błędnej decyzji systemu ML, kwestia odpowiedzialności prawnej pozostaje nierozwiązana.

4. **Walidacja kliniczna** - wymóg przeprowadzenia badań klinicznych walidujących skuteczność i bezpieczeństwo systemów ML.

5. **Integracja z workflow klinicznym** - trudności w adopcji systemów ML w rzeczywistej praktyce klinicznej.

Regulacje prawne, w tym Medical Device Regulation (MDR) w Unii Europejskiej oraz FDA guidance w Stanach Zjednoczonych, coraz bardziej podkreślają wymóg interpretowalności dla systemów ML wykorzystywanych w diagnostyce i prognozowaniu medycznym. To sprawia, że wyjaśnialna sztuczna inteligencja (XAI) staje się nie opcją, lecz koniecznością dla wdrażania ML w medycynie."""

    doc.add_paragraph(content_2_2)

    # 2.3. Metody uczenia maszynowego
    doc.add_heading('2.3. Metody uczenia maszynowego', level=2)

    doc.add_heading('2.3.1. Regresja logistyczna', level=3)

    content_2_3_1 = """Regresja logistyczna stanowi fundamentalną metodę klasyfikacji binarnej, szeroko stosowaną w medycynie ze względu na swoją interpretowalność i prostotę. Model estymuje prawdopodobieństwo przynależności do klasy pozytywnej poprzez funkcję logistyczną (sigmoid):

P(Y=1|X) = 1 / (1 + exp(-β₀ - β₁X₁ - ... - βₙXₙ))

gdzie β₀, β₁, ..., βₙ to parametry modelu szacowane metodą maksymalizacji wiarygodności.

Główną zaletą regresji logistycznej jest naturalna interpretowalność - współczynniki β_i bezpośrednio wskazują kierunek i siłę związku pomiędzy cechą a wynikiem. Wartość exp(β_i) reprezentuje iloraz szans (odds ratio), intuicyjną miarę ryzyka w medycynie.

Ograniczenia regresji logistycznej obejmują założenie liniowości związków, trudności w modelowaniu złożonych interakcji oraz ograniczoną zdolność do aproksymacji nielinearnych granic decyzyjnych."""

    doc.add_paragraph(content_2_3_1)

    # 2.3.2 Random Forest
    doc.add_heading('2.3.2. Lasy losowe (Random Forest)', level=3)

    content_2_3_2 = """Random Forest to metoda ensemble learning oparta na agregacji wielu drzew decyzyjnych. Algorytm wprowadzony przez Breimana (2001) łączy predykcje z wielu drzew trenowanych na losowych podpróbkach danych (bagging) z losowym wyborem cech w każdym węźle rozgałęzienia, co zwiększa różnorodność drzew i redukuje overfitting.

Kluczowe cechy Random Forest:
- **Wysoka dokładność** - agregacja wielu modeli redukuje wariancję predykcji
- **Robustness** - odporność na outliers i missing values
- **Feature importance** - naturalna miara ważności cech
- **Nielinearność** - zdolność modelowania złożonych zależności
- **Minimalna potrzeba tuningu** - działa dobrze z domyślnymi parametrami

W kontekście medycznym, Random Forest wykazuje dobre właściwości dla danych klinicznych charakteryzujących się stosunkowo niewielką liczbą próbek, mieszanymi typami cech oraz potencjalnie złożonymi interakcjami. Caruana i in. (2015) wykazali wysoką skuteczność Random Forest w predykcji ryzyka w szpitalach."""

    doc.add_paragraph(content_2_3_2)

    # 2.3.3 XGBoost
    doc.add_heading('2.3.3. Gradient Boosting (XGBoost, LightGBM)', level=3)

    content_2_3_3 = """Gradient Boosting to rodzina algorytmów ensemble, które sekwencyjnie trenują słabe klasyfikatory (zwykle drzewa decyzyjne), przy czym każdy kolejny model stara się korygować błędy popełnione przez poprzednie modele. XGBoost (eXtreme Gradient Boosting) i LightGBM to zaawansowane implementacje gradient boosting, optymalizowane pod kątem wydajności i dokładności.

**XGBoost**, wprowadzony przez Chen i Guestrin (2016), wykorzystuje regularyzację (L1 i L2) w funkcji straty, co zapobiega overfittingowi. Kluczowe innowacje XGBoost obejmują:
- Approximate algorithm dla znajdowania punktów rozgałęzienia
- Sparse-aware split finding dla obsługi missing values
- Parallel processing i cache optimization
- Built-in cross-validation

**LightGBM**, opracowany przez Microsoft Research, wprowadza alternatywną strategię budowy drzew (leaf-wise zamiast level-wise), co przyspiesza trenowanie przy zachowaniu lub zwiększeniu dokładności. Dodatkowo, LightGBM wykorzystuje Gradient-based One-Side Sampling (GOSS) oraz Exclusive Feature Bundling (EFB) do redukcji złożoności obliczeniowej.

W kontekście medycznym, modele gradient boosting wykazały wysoką skuteczność w licznych zastosowaniach:
- Predykcja śmiertelności w ostrych stanach (sepsis, zawał serca)
- Klasyfikacja nowotworów
- Prognozowanie przebiegu chorób przewlekłych
- Stratyfikacja ryzyka w kardiologii

Szczególnie istotna jest możliwość obsługi niezbalansowania klas poprzez parametr scale_pos_weight (XGBoost) lub is_unbalance (LightGBM), co jest kluczowe dla danych medycznych, gdzie rzadsze zdarzenia (np. zgony) są często bardziej istotne klinicznie."""

    doc.add_paragraph(content_2_3_3)

    # 2.4. Wyjaśnialna sztuczna inteligencja
    doc.add_heading('2.4. Wyjaśnialna sztuczna inteligencja (XAI)', level=2)

    content_2_4 = """Wyjaśnialna sztuczna inteligencja (Explainable Artificial Intelligence, XAI) to dziedzina AI skupiająca się na tworzeniu modeli i metod, które umożliwiają zrozumienie, dlaczego system AI dokonał określonej decyzji lub predykcji. W kontekście medycznym, XAI ma szczególne znaczenie ze względu na:

1. **Wymogi regulacyjne** - MDR i FDA wymagają transparentności algorytmów medycznych
2. **Zaufanie kliniczne** - lekarze muszą rozumieć podstawy decyzji przed zastosowaniem w praktyce
3. **Walidacja merytoryczna** - możliwość weryfikacji, czy model "myśli" zgodnie z wiedzą medyczną
4. **Odkrywanie wiedzy** - identyfikacja nowych biomarkerów i zależności klinicznych
5. **Komunikacja z pacjentem** - obowiązek informed consent wymaga wyjaśnienia podstaw prognoz
6. **Debugging i improvement** - rozumienie błędów modelu pozwala na jego poprawę

Adler-Milstein i Aggarwal (2019) zidentyfikowali pięć kluczowych wymiarów interpretowalności w systemach ML medycznych:
- Transparentność algorytmu (algorithm transparency)
- Zrozumiałość predykcji (prediction understandability)
- Wiarygodność wyjaśnień (explanation fidelity)
- Użyteczność kliniczna (clinical actionability)
- Sprawiedliwość (fairness)

Taksonomia metod XAI obejmuje kilka istotnych dychotomii:

**Model-agnostic vs Model-specific:**
- Model-agnostic (LIME, SHAP, DALEX) działają z dowolnym modelem ML
- Model-specific (Linear/Logistic coefficients, Tree paths, attention weights) wykorzystują wewnętrzną strukturę konkretnego typu modelu

**Local vs Global explanations:**
- Local explanations wyjaśniają pojedynczą predykcję dla konkretnej instancji
- Global explanations opisują ogólne zachowanie modelu na całym zbiorze danych

**Post-hoc vs Inherently interpretable:**
- Post-hoc methods (LIME, SHAP) wyjaśniają istniejące "czarne skrzynki"
- Inherently interpretable models (GAM, EBM, Decision Trees) są transparentne z natury

W niniejszej pracy wykorzystano zarówno metody post-hoc (LIME, SHAP, DALEX) jak i inherently interpretable (EBM), co pozwoliło na kompleksową ocenę interpretowalności systemu predykcyjnego."""

    doc.add_paragraph(content_2_4)

    # 2.5. Metody XAI - szczegółowy przegląd
    doc.add_heading('2.5. Metody XAI - szczegółowy przegląd', level=2)

    # 2.5.1. LIME
    doc.add_heading('2.5.1. LIME (Local Interpretable Model-agnostic Explanations)', level=3)

    content_lime = """LIME (Local Interpretable Model-agnostic Explanations), wprowadzona przez Ribeiro i in. (2016), to metoda post-hoc XAI generująca lokalne wyjaśnienia dla pojedynczych predykcji dowolnego modelu klasyfikacji lub regresji.

**Podstawowa idea LIME:**
Dla danej instancji x, która ma być wyjaśniona, LIME aproksymuje złożony model f lokalnie prostym, interpretowalnym modelem g (np. regresja liniowa). Aproksymacja jest dobra w bliskim sąsiedztwie x, ale nie musi być dokładna globalnie.

**Algorytm LIME:**
1. Generowanie sąsiedztwa: Tworzenie perturbowanych wersji instancji x poprzez losowe modyfikacje wartości cech
2. Ewaluacja modelu: Uzyskanie predykcji modelu f dla wszystkich perturbowanych instancji
3. Ważenie instancji: Przypisanie wag bazujących na odległości od oryginalnej instancji x (bliższe instancje mają większą wagę)
4. Trenowanie interpretowalnego modelu: Dopasowanie prostego modelu liniowego g do przewidywań f w sąsiedztwie x
5. Ekstrakcja wyjaśnienia: Współczynniki modelu g wskazują ważność każdej cechy

Funkcja celu LIME:

ξ(x) = argmin_{g∈G} L(f, g, π_x) + Ω(g)

gdzie:
- L(f, g, π_x) to funkcja straty mierząca jak dobrze g aproksymuje f w sąsiedztwie x zdefiniowanym przez π_x
- Ω(g) to miara złożoności modelu g (regularyzacja)
- G to klasa interpretowalnych modeli (np. modele liniowe)

**Zalety LIME:**
- Model-agnostic - działa z dowolnym klasyfikatorem
- Lokalność - wyjaśnienia dostosowane do konkretnej instancji
- Interpretowalność - wykorzystanie prostych modeli liniowych
- Elastyczność - możliwość kontroli złożoności wyjaśnienia

**Ograniczenia LIME:**
- Niestabilność - wyjaśnienia mogą się różnić pomiędzy uruchomieniami ze względu na losowość
- Wybór sąsiedztwa - wyniki zależą od parametrów generowania perturbacji
- Czas obliczeń - konieczność ewaluacji modelu dla wielu perturbowanych instancji
- Lokalna aproksymacja - może nie odzwierciedlać globalnego zachowania modelu

W kontekście medycznym, LIME stosowano do wyjaśniania predykcji w diagnostyce nowotworów, prognozowaniu sepsy, oraz klasyfikacji obrazów medycznych."""

    doc.add_paragraph(content_lime)

    # 2.5.2. SHAP
    doc.add_heading('2.5.2. SHAP (SHapley Additive exPlanations)', level=3)

    content_shap = """SHAP (SHapley Additive exPlanations), opracowana przez Lundberg i Lee (2017), to ujednolicone podejście do wyjaśniania predykcji modeli ML oparte na teorii gier kooperacyjnych, konkretnie na wartościach Shapleya.

**Teoria wartości Shapleya:**
Wartość Shapleya, wprowadzona przez Lloyda Shapleya w 1953 roku, jest konceptem z teorii gier kooperacyjnych, który przydziela sprawiedliwy udział każdemu graczowi w koalicyjnej grze. W kontekście ML, "gracze" to cechy, a "gra" to predykcja modelu.

Wartość Shapleya dla cechy i definiowana jest jako:

φ_i = Σ_{S ⊆ N \\{i}} [|S|!(|N|-|S|-1)! / |N|!] [f_S∪{i}(x_S∪{i}) - f_S(x_S)]

gdzie:
- N to zbiór wszystkich cech
- S to podzbiór cech nie zawierający cechy i
- f_S(x_S) to wartość predykcji modelu trenowanego tylko na podzbiorze cech S

Intuicyjnie, φ_i reprezentuje średni marginalny wkład cechy i do predykcji, uśredniony po wszystkich możliwych permutacjach cech.

**Właściwości wartości Shapleya:**
1. **Local accuracy:** Σφ_i = f(x) - E[f(X)] (wyjaśnienia sumują się do różnicy predykcji i wartości bazowej)
2. **Missingness:** φ_i = 0 jeśli cecha i nie występuje w modelu
3. **Consistency:** Jeśli zmiana modelu zwiększa marginalny wkład cechy, jej wartość Shapleya nie może zmaleć

**Warianty SHAP:**

1. **TreeSHAP** - zoptymalizowany algorytm dla modeli drzewiastych (RF, XGBoost, LightGBM), redukujący złożoność obliczeniową z wykładniczej do wielomianowej.

2. **KernelSHAP** - aproksymacja wartości Shapleya dla dowolnych modeli, oparta na ważonej regresji liniowej.

3. **LinearSHAP** - dokładne wartości Shapleya dla modeli liniowych, obliczane analitycznie.

4. **DeepSHAP** - adaptacja dla głębokich sieci neuronowych, aproksymująca wartości Shapleya przez backpropagation.

**Zalety SHAP:**
- Solidne podstawy teoretyczne (teoria wartości Shapleya)
- Gwarantowane właściwości matematyczne (local accuracy, consistency, missingness)
- TreeSHAP oferuje dużą prędkość dla modeli drzewiastych
- Wizualizacje (waterfall, force, beeswarm plots) intuicyjne dla klinicystów
- Możliwość zarówno local jak i global explanations

**Ograniczenia SHAP:**
- Czas obliczeń dla KernelSHAP może być znaczący
- Interpretacja wartości Shapleya wymaga zrozumienia koncepcji wartości bazowej
- Założenie niezależności cech (może być problemowe gdy cechy są skorelowane)

W medycynie, SHAP wykorzystano szeroko do wyjaśniania predykcji w sepsiie, COVID-19, onkologii oraz kardiologii, uzyskując wysoką akceptację ze strony klinicystów."""

    doc.add_paragraph(content_shap)

    # 2.5.3 DALEX
    doc.add_heading('2.5.3. DALEX (Descriptive mAchine Learning EXplanations)', level=3)

    content_dalex = """DALEX (Descriptive mAchine Learning EXplanations) to framework opracowany przez Biecka i Burzykowskiego (2021), oferujący zestaw metod do eksploracji, wizualizacji i wyjaśniania modeli uczenia maszynowego. DALEX przyjmuje podejście model-agnostic i skupia się na różnorodnych aspektach interpretowalności.

**Główne komponenty DALEX:**

1. **Model Performance** - ewaluacja ogólnej wydajności modelu z wykorzystaniem różnych metryk

2. **Variable Importance (Permutation Feature Importance)** - globalna ocena ważności cech przez pomiar degradacji wydajności modelu po permutacji wartości danej cechy

3. **Partial Dependence Profiles (PDP)** - wizualizacja marginalnego efektu cechy na predykcję modelu

4. **Individual Conditional Expectation (ICE)** - personalizowane profile pokazujące jak zmiana cechy wpływa na predykcję dla konkretnej instancji

5. **Break-down Analysis** - dekompozycja predykcji na wkład poszczególnych cech (podobne do SHAP, ale z innym podejściem obliczeniowym)

6. **SHAP-like values** - implementacja wartości podobnych do Shapleya

7. **Residual Analysis** - analiza błędów modelu dla identyfikacji systematycznych problemów

DALEX wyróżnia się szczególnie w porównaniach modeli oraz identyfikacji przypadków, dla których model działa nietypowo. Framework jest dostępny zarówno w R jak i Python, z bogatym zestawem wizualizacji.

**Zalety DALEX:**
- Kompleksowe podejście do interpretowalności
- Bogaty zestaw wizualizacji i narzędzi diagnostycznych
- Model-agnostic
- Dobrze rozwinięta dokumentacja i społeczność

**Ograniczenia DALEX:**
- Break-down analysis może być niestabilna (zależna od kolejności cech)
- Permutation importance może być problematyczna dla skorelowanych cech
- Czas obliczeń dla niektórych metod (np. PDP dla wielu cech)"""

    doc.add_paragraph(content_dalex)

    # 2.5.4. EBM
    doc.add_heading('2.5.4. EBM (Explainable Boosting Machine)', level=3)

    content_ebm = """Explainable Boosting Machine (EBM), opracowany przez Nori i in. (2019) w ramach biblioteki InterpretML, to inherently interpretable model łączący wysoką dokładność z pełną transparentnością.

EBM bazuje na Generalized Additive Models (GAM):

g(E[y]) = β₀ + f₁(x₁) + f₂(x₂) + ... + fₙ(xₙ)

gdzie każde f_i jest dowolną funkcją (nie koniecznie liniową) pojedynczej cechy. W EBM, funkcje f_i są uczone przy użyciu gradient boosting, ale z ograniczeniem do małej głębokości drzew i iteracyjnego trenowania (round-robin), co zachowuje addytywną strukturę.

**Kluczowe innowacje EBM:**

1. **Bagged gradient boosting** - zastosowanie baggingu do redukcji wariancji typowej dla GB

2. **Feature binning** - dyskretyzacja cech ciągłych dla zwiększenia interpretowalności i redukcji overfittingu

3. **Interakcje parami** - możliwość włączenia wybranych interakcji drugiego rzędu (f_ij(x_i, x_j)) przy zachowaniu interpretowalności

4. **Learning rate schedules** - adaptacyjny learning rate dla stabilizacji trenowania

**Zalety EBM:**
- Inherent interpretability - model jest transparentny z natury, nie wymaga post-hoc explanations
- Wizualizacja efektów cech - możliwość bezpośredniej wizualizacji f_i(x_i)
- Accuracy porównywalna z XGBoost/Random Forest
- Wykrywanie interakcji - automatyczne wykrywanie istotnych interakcji
- Monotonicity constraints - możliwość wymuszenia monotoniczności zgodnie z wiedzą domenową

**Ograniczenia EBM:**
- Czas trenowania dłuższy niż standardowy GB
- Ograniczone do addytywnych modeli (interakcje wyższych rzędów trudne)
- Mniej elastyczny niż pełne drzewa decyzyjne

W medycynie, EBM zastosowano do predykcji ryzyka w pneumonii, readmisji szpitalnych oraz sepsiie, uzyskując dokładność konkurencyjną wobec "czarnych skrzynek" przy pełnej transparentności."""

    doc.add_paragraph(content_ebm)

    # 2.6. Metryki medyczne
    doc.add_heading('2.6. Metryki ewaluacji w kontekście medycznym', level=2)

    content_2_6 = """W medycynie, wybór metryk ewaluacji modelu predykcyjnego ma kluczowe znaczenie i różni się od typowych zastosowań ML. Accuracy (dokładność) jest często niewystarczającą metryką ze względu na niezbalansowanie klas, różne koszty błędów typu I i II oraz specyficzne wymagania kliniczne.

**Czułość (Sensitivity, Recall, True Positive Rate):**

Sensitivity = TP / (TP + FN)

Czułość mierzy proporcję prawidłowo zidentyfikowanych przypadków pozytywnych (np. pacjentów z ryzykiem zgonu). W kontekście predykcji śmiertelności, wysoka czułość jest krytyczna - przeoczenie pacjenta wysokiego ryzyka (false negative) może mieć tragiczne konsekwencje. Dlatego w systemach medycznych często przyjmuje się próg sensitivity ≥ 0.80-0.90.

**Swoistość (Specificity, True Negative Rate):**

Specificity = TN / (TN + FP)

Swoistość mierzy proporcję prawidłowo zidentyfikowanych przypadków negatywnych. Niska swoistość prowadzi do nadmiernej liczby fałszywych alarmów, co może skutkować nadmiernym leczeniem, niepotrzebnym stresem pacjentów oraz marnotrawstwem zasobów.

**PPV (Positive Predictive Value, Precision):**

PPV = TP / (TP + FP)

PPV określa prawdopodobieństwo, że pacjent z pozytywną predykcją rzeczywiście jest w grupie ryzyka. W populacjach o niskiej prevalence (częstości występowania), nawet testy o wysokiej czułości i swoistości mogą mieć niskie PPV.

**NPV (Negative Predictive Value):**

NPV = TN / (TN + FN)

NPV określa prawdopodobieństwo, że pacjent z negatywną predykcją rzeczywiście nie jest w grupie ryzyka. Wysokie NPV (≥0.90) jest często wymagane do wykluczenia diagnozy.

**AUC-ROC (Area Under the Receiver Operating Characteristic Curve):**
AUC-ROC mierzy zdolność modelu do dyskryminacji pomiędzy klasami w całym zakresie progów decyzyjnych. Wartość 0.5 odpowiada losowemu klasyfikatorowi, 1.0 - idealnemu. W medycynie, AUC-ROC ≥ 0.75 uznawane jest za akceptowalne, ≥ 0.85 za dobre, a ≥ 0.95 za doskonałe.

**Brier Score:**
Brier Score mierzy kalibrację prawdopodobieństw predykcyjnych:

BS = (1/N) Σ (p_i - y_i)²

gdzie p_i to predykowane prawdopodobieństwo, a y_i to rzeczywisty wynik (0 lub 1). Niższy Brier Score oznacza lepszą kalibrację.

**Balanced Accuracy:**
Balanced Accuracy = (Sensitivity + Specificity) / 2

Szczególnie użyteczna dla niezbalansowanych zbiorów danych.

**Matthews Correlation Coefficient (MCC):**
MCC uwzględnia wszystkie elementy macierzy pomyłek i jest uznawana za jedną z najbardziej zrównoważonych metryk dla klasyfikacji binarnej, szczególnie w przypadku niezbalansowania.

W kontekście medycznym, zaleca się raportowanie wielu metryk jednocześnie, z priorytetem dla czułości i NPV w przypadku wykluczania poważnych stanów (rule-out) oraz swoistości i PPV w przypadku potwierdzania diagnoz (rule-in)."""

    doc.add_paragraph(content_2_6)

    # 2.7. Systemy wsparcia decyzji klinicznych
    doc.add_heading('2.7. Systemy wsparcia decyzji klinicznych', level=2)

    content_2_7 = """Systemy wsparcia decyzji klinicznych (Clinical Decision Support Systems, CDSS) to aplikacje informatyczne zaprojektowane do wspomagania lekarzy i innych pracowników służby zdrowia w podejmowaniu decyzji diagnostycznych i terapeutycznych. CDSS oparte na ML/AI stanowią najnowszą generację takich systemów.

**Architektura CDSS:**
Nowoczesne CDSS oparte na ML zazwyczaj składają się z:
1. Moduł pozyskiwania danych (integracja z EHR/EMR)
2. Moduł przetwarzania i feature engineering
3. Moduł predykcyjny (model ML)
4. Moduł wyjaśnień (XAI)
5. Moduł prezentacji wyników (interfejs użytkownika)
6. Moduł monitorowania i feedbacku

**Integracja z systemami szpitalnymi:**
Kluczowe wyzwania obejmują:
- Standaryzację danych (HL7, FHIR)
- Real-time processing
- Fault tolerance i availability
- Privacy i security (HIPAA, GDPR)
- Weryfikacja i walidacja medyczna

**Interfejsy użytkownika:**
Projektowanie interfejsów dla CDSS wymaga uwzględnienia:
- Workflow kliniczny - minimalny wpływ na czas pracy lekarza
- Clarity and actionability - wyjaśnienia muszą być zrozumiałe i actionable
- Trust and transparency - budowanie zaufania przez transparentność
- Różne grupy użytkowników - dostosowanie do lekarzy, pielęgniarek, pacjentów

**Bezpieczeństwo i guardrails:**
Systemy CDSS muszą zawierać mechanizmy zabezpieczające:
- Wykrywanie sytuacji kryzysowych (np. myśli samobójcze)
- Disclaimery medyczne
- Ograniczenia zakresu działania (np. brak diagnozowania)
- Logowanie i audit trail
- Mechanizmy override dla lekarzy

Sutton i in. (2020) wykazali, że CDSS z odpowiednimi guardrails są bezpieczniejsze i częściej akceptowane przez klinicystów."""

    doc.add_paragraph(content_2_7)

    # 2.8. Podsumowanie
    doc.add_heading('2.8. Podsumowanie i luki w badaniach', level=2)

    content_2_8 = """Przegląd literatury wykazał znaczący postęp w zakresie uczenia maszynowego w medycynie oraz rozwój metod wyjaśnialnej sztucznej inteligencji. Jednakże zidentyfikowano następujące luki badawcze:

1. **Brak systematycznych porównań metod XAI** w kontekście konkretnych zastosowań medycznych - większość prac stosuje pojedynczą metodę XAI.

2. **Ograniczona adaptacja komunikacji** - niewiele systemów dostosowuje formę i złożoność wyjaśnień do poziomu wiedzy medycznej odbiorcy (lekarz vs pacjent).

3. **Niewystarczająca uwaga na guardrails** - aspekty bezpieczeństwa medycznego i etycznego często pomijane w pracach badawczych.

4. **Rzadkie zastosowania XAI w chorobach rzadkich** - większość badań koncentruje się na powszechnych schorzeniach (rak, choroby serca), zaniedbując rzadsze jednostki chorobowe jak zapalenie naczyń.

5. **Brak integracji LLM z XAI** - potencjał dużych modeli językowych do generowania naturalnych wyjaśnień w oparciu o output metod XAI pozostaje niewykorzystany.

Niniejsza praca adresuje powyższe luki poprzez:
- Systematyczne porównanie czterech metod XAI (LIME, SHAP, DALEX, EBM)
- Implementację komunikacji wielopoziomowej (clinician/patient perspectives)
- Kompleksowy system guardrails
- Aplikację do zapalenia naczyń (choroba rzadka)
- Integrację LLM (RAG) dla generowania wyjaśnień naturalnych

Tym samym, praca wnosi oryginalny wkład zarówno w metodykę XAI jak i w praktyczne zastosowania ML w reumatologii."""

    doc.add_paragraph(content_2_8)

    doc.add_page_break()


# Kontynuacja w kolejnych funkcjach...
